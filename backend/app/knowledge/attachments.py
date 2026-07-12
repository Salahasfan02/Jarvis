"""Conversation attachments: files and images uploaded into a chat.

Text is extracted immediately (PDF, Word, Excel, CSV, JSON, code, zip
listings...). Images get Apple-Vision OCR plus, when a vision model is
configured, a natural-language description — so ANY chat model can reason
about them. Attachments persist with the conversation until removed.
"""
from __future__ import annotations

import base64
import io
import tempfile
import zipfile
from pathlib import Path

from .. import db
from ..config import settings
from .store import extract_text as extract_plain

IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff", ".heic"}
MAX_TEXT = 20000


def _extract_docx(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i > 300:
                parts.append("… (truncated)")
                break
            parts.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(parts)


def _extract_zip(data: bytes) -> str:
    parts = ["ZIP archive contents:"]
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for info in zf.infolist()[:200]:
            parts.append(f"- {info.filename} ({info.file_size} bytes)")
        # inline small text files
        text_members = [i for i in zf.infolist()
                        if not i.is_dir() and i.file_size < 40000 and
                        Path(i.filename).suffix.lower() in
                        {".txt", ".md", ".py", ".js", ".ts", ".json", ".csv",
                         ".html", ".css", ".yml", ".yaml", ".toml", ".sh"}][:12]
        for member in text_members:
            content = zf.read(member).decode("utf-8", errors="replace")
            parts.append(f"\n### {member.filename}\n{content[:3000]}")
    return "\n".join(parts)


async def _describe_image(data: bytes, name: str) -> str:
    parts = []
    # exact text via Apple Vision OCR
    try:
        from ..vision import ocr
        suffix = Path(name).suffix or ".png"
        tmp = Path(tempfile.mkstemp(suffix=suffix)[1])
        tmp.write_bytes(data)
        try:
            text = await ocr.ocr_file(tmp)
            if text.strip():
                parts.append("Text found in the image (OCR):\n" + text[:4000])
        finally:
            tmp.unlink(missing_ok=True)
    except Exception:
        pass
    # visual description via the configured multimodal model
    vision_model = settings.get("ollama.vision_model", "")
    if vision_model:
        try:
            from ..llm import ollama_client
            description = await ollama_client.chat_once(
                [{"role": "user",
                  "content": "Describe this image thoroughly: contents, layout, "
                             "any diagrams or UI elements, and anything notable.",
                  "images": [base64.b64encode(data).decode()]}],
                model=vision_model)
            if description:
                parts.append("Visual description:\n" + description[:3000])
        except Exception:
            pass
    return "\n\n".join(parts) or "(image uploaded — no text or description extracted)"


async def ingest(conv_id: str, name: str, data: bytes) -> dict:
    suffix = Path(name).suffix.lower()
    if suffix in IMAGE_EXT:
        text = await _describe_image(data, name)
        return db.add_attachment(conv_id, name, "image", text[:MAX_TEXT],
                                 image_b64=base64.b64encode(data).decode())
    if suffix == ".docx":
        text = _extract_docx(data)
    elif suffix in (".xlsx", ".xlsm"):
        text = _extract_xlsx(data)
    elif suffix == ".zip":
        text = _extract_zip(data)
    else:  # pdf / text / code / config / logs — handled by the knowledge extractor
        text = extract_plain(name, data)
    return db.add_attachment(conv_id, name, "file", text[:MAX_TEXT])


def prompt_block(conv_id: str) -> str:
    """Attachment contents injected into the conversation prompt."""
    attachments = db.list_attachments(conv_id, with_content=True)
    if not attachments:
        return ""
    budget = 24000
    parts = ["The user has attached these files to this conversation:"]
    for att in attachments:
        chunk = att["content_text"][: max(1500, budget // len(attachments))]
        parts.append(f"\n=== ATTACHMENT: {att['name']} ({att['kind']}) ===\n{chunk}")
    return "\n".join(parts)[:26000]
